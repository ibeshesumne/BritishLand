from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from jinja2 import Environment, FileSystemLoader

def set_paragraph_spacing(paragraph, space_before=Pt(0), space_after=Pt(0)):
    """Function to set spacing before and after a paragraph."""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = space_before
    paragraph_format.space_after = space_after

# Create a new Document
doc = Document()

# Add a title with two rows
title = doc.add_heading('', 0)
run = title.add_run('Jake Leary')
run.font.size = Pt(24)
run = title.add_run('\nFlat 25, Voltaire Building, 330 Garratt Lane, SW18 4FQ')
run.font.size = Pt(12)
run = title.add_run('\n+44 7454 606 300 · jakeleary@gmail.com')
run.font.size = Pt(12)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Summary section
summary_heading = doc.add_paragraph('Summary')
set_paragraph_spacing(summary_heading, space_after=Pt(6))  # Adjust spacing as needed
summary_heading.runs[0].bold = True
summary = doc.add_paragraph('As a highly motivated economics graduate, I am passionate about using economic principles to tackle complex societal issues. My studies helped me understand key economic theory and I am keen to apply this knowledge to real-world problems. My knowledge of behavioural economics was useful for the analysis of data relevant to the profile and decisions made by consumers.\n\nIn my current role, assisting with the day-to-day financial accounts taught me the importance of optimising efficiency in the long term interests of the company. It enhanced my ability to make a useful contribution in the analysis of data for the team involved in the critical policy decision process.\n\nI am dedicated to leveraging the skills acquired through my education and current role to actively contribute to sustainable development, striving for a career that not only aligns with my passion but also has a positive impact on society.')
summary.paragraph_format.left_indent = Inches(0.5)

# Skills
skills_heading = doc.add_paragraph('Skills')
set_paragraph_spacing(skills_heading, space_after=Pt(4)) 
skills_heading.runs[0].bold = True
skills = doc.add_paragraph()
skills.add_run('Languages spoken: English (Native); Japanese (Successfully passed N4 Japanese government proficiency test; French (completed A level French)\n\n')
skills.add_run('Python: currently enrolled in Harvard University\'s CS50 python course on artificial intelligence; Python libraries used in past include: NumPy, Pandas, Matplotlib, NLTK\n\n')
skills.add_run('Stata: used for econometric analysis; regression; interpreting results using individual fixed effects, pooled ordinalry least squares\n\n')
skills.add_run('Microsoft excel: used for econometric analysis; regression; interpreting results using individual fixed effects, pooled ordinary least squares\n\n')
skills.paragraph_format.left_indent = Inches(0.5)

# Experience
experience_heading = doc.add_paragraph('Professional Experience')
set_paragraph_spacing(experience_heading, space_after=Pt(0)) 
experience_heading.runs[0].bold = True
experience = doc.add_paragraph()
run = experience.add_run('Finance Assistant, Love Cocoa & H!P Chocolate')
run.bold = True
experience.add_run('\nApril 2023 - Present\n')
experience.add_run('\nData Analysis: Conducted weekly analysis of grocery sales data for our H!P chocolate sales, utilising Excel and data visualisation tools. Provided actionable insights for strategic decision-making for the marketing and sales team.\n\nFinancial Reporting: Generated and presented weekly reports on debt recovery, total invoices and credits being made, allowing for a clearer image of company performance. Assisted with end-of-month reports for senior leadership. Supported the finance director in financial forecasts for both brands.\n\nDebt recovery: Through active monitoring of accounts receivable I successfully recovered £500,000 in overdue invoices through debt collection efforts therefore improving the company\'s overall cash flow and financial stability. Maintained positive client relationships while ensuring timely payments.\n\nProcess optimization: Collaborated with the operations team to streamline invoicing processes, reduce errors and reduce payment delays due to credit requests. Implementing brand coding of all revenue allowed for a clearer image of brand performance in the profit and loss statement.\n\n')
experience.paragraph_format.left_indent = Inches(0.5)
experience = doc.add_paragraph()
run = experience.add_run('Supervisor - Front of house, The Dynamo')
run.bold = True
experience.add_run('\nOctober 2021 - October 2022\n')
experience.add_run('\nWaiter and bartender skills: It was a challenge, a pleasure and a life learning experience to interact with the staff and the customers who represented a wide cross section of international backgrounds. Mentored and trained new team members.\n\n')
experience.paragraph_format.left_indent = Inches(0.5)

# Adjust spacing after the second professional experience
set_paragraph_spacing(experience, space_after=Pt(1)) 

# Education
education_heading = doc.add_paragraph('Education')
set_paragraph_spacing(education_heading, space_after=Pt(2))  
education_heading.runs[0].bold = True
education = doc.add_paragraph()
run = education.add_run('University of Sussex, Bachelor of Science (Hons): Economics - Result 2:1')
run.bold = True
education.add_run('\n2017-2021\n')
education.add_run('\nModules of interest included:\n\n')
bold_run = education.add_run('Behavioural Economics')
bold_run.bold = True
education.add_run(': analysis of irrational decisions which leads to a decline in efficiency in certain business decisions.\n\n')
bold_run = education.add_run('Financial Economics')
bold_run.bold = True
education.add_run(': The module required the use of excel to perform financial analysis such as use of the capital asset price model. \n\n')
bold_run = education.add_run('Big Data and Economics')
bold_run.bold = True
education.add_run(': The module focused on the use of python to perform machine learning techniques and data analysis.\n\n')
bold_run = education.add_run('Applied Econometrics')
bold_run.bold = True
education.add_run(': This module used Stata to perform regression analysis to interpret results from a dataset.\n\n')
education.paragraph_format.left_indent = Inches(0.5)

education = doc.add_paragraph()
run = education.add_run('Millfield School, UK')
run.bold = True
education.add_run('\nSeptember 2013 - June 2017\n')
education.add_run('\nA-Levels: Economics, Japanese, French, Mathematics; 10 GCSEs\n\n ')
education.paragraph_format.left_indent = Inches(0.5)

# Interests
interests_heading = doc.add_paragraph('Interests')
set_paragraph_spacing(interests_heading, space_after=Pt(2)) 
interests_heading.runs[0].bold = True
interests = doc.add_paragraph()
run = interests.add_run('Trekking in Nepal, travelling in Vietnam and Australia, walking in Switzerland. Discovery of cycling.')
run.bold = False
interests.paragraph_format.left_indent = Inches(0.5)

# Save the document
doc.save('cv.docx')

# Generate HTML CV
env = Environment(loader=FileSystemLoader('.'))
template = env.get_template('cv_template.html')

# Extract content from the DOCX file
content = {
    'title': 'Jake Leary',
    'address': 'Flat 25, Voltaire Building, 330 Garratt Lane, SW18 4FQ',
    'contact': '+44 7454 606 300 · jakeleary@gmail.com',
    'summary': 'As a highly motivated economics graduate, I am passionate about using economic principles to tackle complex societal issues. My studies helped me understand key economic theory and I am keen to apply this knowledge to real-world problems. My knowledge of behavioural economics was useful for the analysis of data relevant to the profile and decisions made by consumers.\n\nIn my current role, assisting with the day-to-day financial accounts taught me the importance of optimising efficiency in the long term interests of the company. It enhanced my ability to make a useful contribution in the analysis of data for the team involved in the critical policy decision process.\n\nI am dedicated to leveraging the skills acquired through my education and current role to actively contribute to sustainable development, striving for a career that not only aligns with my passion but also has a positive impact on society.',
    'skills': [
        'Languages spoken: English (Native); Japanese (Successfully passed N4 Japanese government proficiency test; French (completed A level French)',
        'Python: currently enrolled in Harvard University\'s CS50 python course on artificial intelligence; Python libraries used in past include: NumPy, Pandas, Matplotlib, NLTK',
        'Stata: used for econometric analysis; regression; interpreting results using individual fixed effects, pooled ordinalry least squares',
        'Microsoft excel: used for econometric analysis; regression; interpreting results using individual fixed effects, pooled ordinary least squares'
    ],
    'experience': [
        {
            'title': 'Finance Assistant, Love Cocoa & H!P Chocolate',
            'date': 'April 2023 - Present',
            'description': [
                'Data Analysis: Conducted weekly analysis of grocery sales data for our H!P chocolate sales, utilising Excel and data visualisation tools. Provided actionable insights for strategic decision-making for the marketing and sales team.',
                'Financial Reporting: Generated and presented weekly reports on debt recovery, total invoices and credits being made, allowing for a clearer image of company performance. Assisted with end-of-month reports for senior leadership. Supported the finance director in financial forecasts for both brands.',
                'Debt recovery: Through active monitoring of accounts receivable I successfully recovered £500,000 in overdue invoices through debt collection efforts therefore improving the company\'s overall cash flow and financial stability. Maintained positive client relationships while ensuring timely payments.',
                'Process optimization: Collaborated with the operations team to streamline invoicing processes, reduce errors and reduce payment delays due to credit requests. Implementing brand coding of all revenue allowed for a clearer image of brand performance in the profit and loss statement.'
            ]
        },
        {
            'title': 'Supervisor - Front of house, The Dynamo',
            'date': 'October 2021 - October 2022',
            'description': [
                'Waiter and bartender skills: It was a challenge, a pleasure and a life learning experience to interact with the staff and the customers who represented a wide cross section of international backgrounds. Mentored and trained new team members.'
            ]
        }
    ],
    'education': [
        {
            'title': 'University of Sussex, Bachelor of Science (Hons): Economics - Result 2:1',
            'date': '2017-2021',
            'description': [
                'Modules of interest included:',
                'Behavioural Economics: analysis of irrational decisions which leads to a decline in efficiency in certain business decisions.',
                'Financial Economics: The module required the use of excel to perform financial analysis such as use of the capital asset price model.',
                'Big Data and Economics: The module focused on the use of python to perform machine learning techniques and data analysis.',
                'Applied Econometrics: This module used Stata to perform regression analysis to interpret results from a dataset.'
            ]
        },
        {
            'title': 'Millfield School, UK',
            'date': 'September 2013 - June 2017',
            'description': [
                'A-Levels: Economics, Japanese, French, Mathematics; 10 GCSEs'
            ]
        }
    ],
    'interests': 'Trekking in Nepal, travelling in Vietnam and Australia, walking in Switzerland. Discovery of cycling.'
}

html_output = template.render(content)

with open('cv.html', 'w') as f:
    f.write(html_output)